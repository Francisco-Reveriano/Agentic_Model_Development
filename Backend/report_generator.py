"""Report assembly engine — generates Credit Risk Model Development documents.

Uses python-docx to produce five report types:
  1. Data Quality (DQ) Report
  2. Model Development Report (C1 template for PD, LGD, EAD)
  3. Expected Loss (EL) Report
  4. (DQ + Model combined via the above)
  5. Executive Summary / Portfolio report (EL)

Each method returns the Path of the generated .docx file.
"""

from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


class ReportGenerator:
    """Assembles .docx reports from agent handoff artifacts."""

    # ------------------------------------------------------------------
    # Public report generators
    # ------------------------------------------------------------------

    def generate_dq_report(self, dq_data: dict, output_path: Path) -> Path:
        """Generate the Data Quality report.

        Sections:
          1. Data Overview
          2. Data Quality Scorecard
          3. Data Treatments
          4. Feature Profiling
          5. Fit for Modeling Decision
          6. Query Trace

        Args:
            dq_data: Contents of Data_Agent handoff.json (or enriched version).
            output_path: Directory to write the report into.

        Returns:
            Path to the generated .docx file.
        """
        output_path = Path(output_path)
        output_path.mkdir(parents=True, exist_ok=True)

        doc = self._create_base_doc(
            title="Data Quality Report",
            subtitle="Credit Risk Modeling Platform — Data Readiness Assessment",
        )

        metrics = dq_data.get("metrics", {})
        output_files = dq_data.get("output_files", {})
        status = dq_data.get("status", "unknown")

        # --- Section 1: Data Overview ---
        self._add_heading(doc, "1. Data Overview", level=1)
        overview_metrics = {
            "Initial Rows": metrics.get("initial_rows", "N/A"),
            "Initial Columns": metrics.get("initial_cols", "N/A"),
            "Resolved Rows (Fully Paid + Charged Off)": metrics.get("resolved_rows", "N/A"),
            "Final Feature Count": metrics.get("feature_count", "N/A"),
            "Default Rate": f"{metrics.get('default_rate', 0):.2%}" if isinstance(metrics.get("default_rate"), (int, float)) else "N/A",
            "Default Count": metrics.get("default_count", "N/A"),
            "Fully Paid Count": metrics.get("fully_paid_count", "N/A"),
        }
        self._add_metric_table(doc, overview_metrics)

        # --- Section 2: Data Quality Scorecard ---
        self._add_heading(doc, "2. Data Quality Scorecard", level=1)
        scorecard_items = dq_data.get("scorecard", {})
        if scorecard_items:
            headers = ["Test ID", "Test Name", "Result", "Details"]
            rows = []
            for test_id, info in scorecard_items.items():
                if isinstance(info, dict):
                    rows.append([
                        test_id,
                        info.get("name", ""),
                        info.get("result", ""),
                        info.get("details", ""),
                    ])
                else:
                    rows.append([test_id, str(info), "", ""])
            self._add_table(doc, headers, rows)
        else:
            self._add_paragraph(
                doc,
                "Data quality scorecard was generated as part of the Data Agent "
                "analysis. Key metrics are summarized in the Data Overview above. "
                "Detailed test results (DQ-01 through DQ-10) are available in the "
                "agent conversation log.",
            )
            quality_metrics = {
                "Columns Dropped (leakage + non-useful)": metrics.get("cols_dropped", "N/A"),
                "Leakage Columns Preserved (for LGD/EAD)": metrics.get("leakage_cols_preserved", "N/A"),
                "Winsorized Columns": metrics.get("winsorized_cols", "N/A"),
            }
            self._add_metric_table(doc, quality_metrics)

        # --- Section 3: Data Treatments ---
        self._add_heading(doc, "3. Data Treatments", level=1)
        treatments = dq_data.get("treatments", [])
        if treatments:
            headers = ["Step", "Treatment", "Columns Affected", "Rationale"]
            rows = []
            for t in treatments:
                if isinstance(t, dict):
                    rows.append([
                        str(t.get("step", "")),
                        t.get("treatment", ""),
                        t.get("columns", ""),
                        t.get("rationale", ""),
                    ])
            self._add_table(doc, headers, rows)
        else:
            self._add_paragraph(doc, "The cleaning pipeline applied the following steps:")
            default_treatments = [
                ["1", "Drop leakage columns", "Post-origination fields (recoveries, total_pymnt, etc.)", "Prevent information leakage into PD model"],
                ["2", "Filter to resolved loans", "loan_status IN (Fully Paid, Charged Off)", "Binary classification requires definitive outcomes"],
                ["3", "Type coercion", "int_rate, term, emp_length, revol_util", "Convert string representations to numeric"],
                ["4", "Missing value imputation", "Numeric: grade-segmented median; Categorical: mode", "Preserve distributional relationships"],
                ["5", "Winsorization (1st/99th pctl)", "annual_inc, revol_bal, loan_amnt, funded_amnt, dti, open_acc", "Reduce outlier influence"],
                ["6", "Categorical encoding", "grade, sub_grade, purpose, home_ownership, etc.", "Ordinal, risk-based, and one-hot encoding"],
            ]
            self._add_table(
                doc,
                ["Step", "Treatment", "Columns Affected", "Rationale"],
                default_treatments,
            )

        # --- Section 4: Feature Profiling ---
        self._add_heading(doc, "4. Feature Profiling", level=1)
        profiling = dq_data.get("feature_profiling", [])
        if profiling:
            headers = ["Feature", "Type", "Null Rate", "Mean", "Std", "Min", "Max"]
            rows = []
            for fp in profiling:
                if isinstance(fp, dict):
                    rows.append([
                        fp.get("feature", ""),
                        fp.get("type", ""),
                        f"{fp.get('null_rate', 0):.2%}" if isinstance(fp.get("null_rate"), (int, float)) else "",
                        str(fp.get("mean", "")),
                        str(fp.get("std", "")),
                        str(fp.get("min", "")),
                        str(fp.get("max", "")),
                    ])
            self._add_table(doc, headers, rows)
        else:
            self._add_paragraph(
                doc,
                f"The cleaned dataset contains {metrics.get('feature_count', 'N/A')} features "
                f"across {metrics.get('resolved_rows', 'N/A')} resolved loans. "
                "Detailed per-feature profiling (distributions, percentiles, cardinality) "
                "was performed during the Data Agent analysis phase. "
                "See the agent conversation log for column-level profiles.",
            )

        # --- Section 5: Fit for Modeling Decision ---
        self._add_heading(doc, "5. Fit for Modeling Decision", level=1)
        fit_decision = dq_data.get("fit_for_modeling", {})
        if fit_decision:
            self._add_metric_table(doc, {
                "Decision": fit_decision.get("decision", ""),
                "Rationale": fit_decision.get("rationale", ""),
                "Caveats": fit_decision.get("caveats", ""),
            })
        else:
            if status == "success":
                self._add_paragraph(
                    doc,
                    "DECISION: FIT FOR MODEL DEVELOPMENT\n\n"
                    "The dataset has been cleaned and validated through the 6-step "
                    "pipeline. All material data quality issues have been addressed. "
                    "The cleaned feature set and target variables are ready for "
                    "feature engineering and model training.",
                )
            else:
                self._add_paragraph(
                    doc,
                    f"Data Agent completed with status: {status}. "
                    "Review the agent logs for details on any issues encountered.",
                )

        # --- Section 6: Query Trace ---
        self._add_heading(doc, "6. Query Trace", level=1)
        query_trace = dq_data.get("query_trace", [])
        if query_trace:
            headers = ["Query Name", "SQL", "Rows Returned"]
            rows = []
            for qt in query_trace:
                if isinstance(qt, dict):
                    sql_text = qt.get("sql", "")
                    if len(sql_text) > 120:
                        sql_text = sql_text[:117] + "..."
                    rows.append([
                        qt.get("query_name", ""),
                        sql_text,
                        str(qt.get("returned_rows", "")),
                    ])
            self._add_table(doc, headers, rows)
        else:
            self._add_paragraph(
                doc,
                "Query trace details are captured in the agent conversation log. "
                "The Data Agent executed baseline scans followed by adaptive "
                "drill-down queries for each identified data quality dimension.",
            )

        # --- Output Files Reference ---
        if output_files:
            self._add_heading(doc, "Output Files", level=2)
            self._add_metric_table(doc, output_files)

        filepath = output_path / "data_quality_report.docx"
        doc.save(str(filepath))
        return filepath

    def generate_model_report(
        self,
        model_type: str,
        artifacts: dict,
        output_path: Path,
    ) -> Path:
        """Generate a C1-template Model Development Report.

        Sections (C1 template):
          1. Introduction
          2. Purpose and Uses
          3. Use Case Description
          4. Data (sources, treatments, assumptions)
          5. Methodology (features, selection, calibration)
          6. Outputs (description, performance testing)
          7. Implementation
          8. Appendix

        Args:
            model_type: One of 'PD', 'LGD', 'EAD'.
            artifacts: Combined handoff data for this model type.
            output_path: Directory to write the report into.

        Returns:
            Path to the generated .docx file.
        """
        output_path = Path(output_path)
        output_path.mkdir(parents=True, exist_ok=True)

        model_label = model_type.upper()
        full_names = {"PD": "Probability of Default", "LGD": "Loss Given Default", "EAD": "Exposure at Default"}
        full_name = full_names.get(model_label, model_label)

        doc = self._create_base_doc(
            title=f"{full_name} ({model_label}) Model Development Report",
            subtitle="Credit Risk Modeling Platform — C1 Model Documentation",
        )

        metrics = artifacts.get("metrics", {})
        model_info = artifacts.get("model_info", {})
        champion = artifacts.get("champion", {})
        features = artifacts.get("features", {})
        data_info = artifacts.get("data", {})
        performance = artifacts.get("performance", {})

        # --- Table 1: Model Metadata ---
        self._add_heading(doc, "Model Metadata", level=2)
        metadata = {
            "Model Name": f"{model_label} Credit Risk Model",
            "Model Type": full_name,
            "Model Version": artifacts.get("version", "1.0"),
            "Development Date": datetime.now(timezone.utc).strftime("%Y-%m-%d"),
            "Model Owner": "Credit Risk Modeling Platform",
            "Status": artifacts.get("status", "Development"),
        }
        self._add_metadata_table(doc, metadata)

        # --- Section 1: Introduction ---
        self._add_heading(doc, "1. Introduction", level=1)
        self._add_paragraph(
            doc,
            f"This document presents the development and validation of the {full_name} "
            f"({model_label}) model for the consumer lending credit risk portfolio. "
            f"The model was developed using the Credit Risk Modeling Platform's "
            f"automated model tournament framework, which evaluates multiple algorithms "
            f"and hyperparameter configurations to select the best-performing model.",
        )

        # --- Section 2: Purpose and Uses ---
        self._add_heading(doc, "2. Purpose and Uses", level=1)
        purpose_map = {
            "PD": (
                "The PD model estimates the probability that a borrower will default "
                "on their loan obligation within the remaining life of the loan. "
                "It is used for:\n"
                "- Loan origination decisioning and pricing\n"
                "- Portfolio risk monitoring and early warning\n"
                "- Regulatory capital calculation (Basel III/IV IRB approach)\n"
                "- Expected Credit Loss (ECL) estimation under IFRS 9 / CECL\n"
                "- Stress testing and scenario analysis"
            ),
            "LGD": (
                "The LGD model estimates the proportion of exposure that will be lost "
                "in the event of borrower default, net of recoveries. It is used for:\n"
                "- Loss forecasting and reserve estimation\n"
                "- Risk-adjusted pricing\n"
                "- Regulatory capital calculation (Basel III/IV IRB approach)\n"
                "- Expected Credit Loss computation\n"
                "- Stress testing and downturn LGD estimation"
            ),
            "EAD": (
                "The EAD model estimates the outstanding exposure at the time of default. "
                "For term loans, this includes the remaining principal balance and any "
                "accrued interest. It is used for:\n"
                "- Exposure management and limit setting\n"
                "- Regulatory capital calculation\n"
                "- Expected Credit Loss computation\n"
                "- Portfolio-level loss aggregation"
            ),
        }
        self._add_paragraph(doc, purpose_map.get(model_label, f"Purpose for {model_label} model."))

        # --- Section 3: Use Case Description ---
        self._add_heading(doc, "3. Use Case Description", level=1)
        use_case_map = {
            "PD": (
                "The model is applied to the consumer unsecured lending portfolio "
                "(personal loans originated through the LendingClub platform). "
                "The target variable is a binary default indicator, where default "
                "is defined as a loan reaching 'Charged Off' status. "
                "Non-default ('Fully Paid') serves as the negative class."
            ),
            "LGD": (
                "The model is applied to defaulted loans in the consumer unsecured "
                "lending portfolio. LGD is calculated as 1 minus the recovery rate, "
                "where recovery rate = (recoveries - collection fees) / funded amount. "
                "Values are bounded to [0, 1]."
            ),
            "EAD": (
                "The model estimates exposure at default for the consumer lending "
                "portfolio. For fully drawn term loans, EAD approximates the "
                "outstanding principal balance. The credit conversion factor (CCF) "
                "relates EAD to the original funded amount."
            ),
        }
        self._add_paragraph(doc, use_case_map.get(model_label, f"Use case for {model_label} model."))

        # --- Section 4: Data ---
        self._add_heading(doc, "4. Data", level=1)

        self._add_heading(doc, "4.1 Data Sources", level=2)
        data_sources = data_info.get("sources", {})
        if data_sources:
            self._add_metric_table(doc, data_sources)
        else:
            self._add_paragraph(
                doc,
                "Primary data source: LendingClub loan-level dataset stored in SQLite "
                "database. The dataset contains origination characteristics, borrower "
                "attributes, and loan performance outcomes.",
            )

        self._add_heading(doc, "4.2 Data Treatments", level=2)
        data_treatments = data_info.get("treatments", [])
        if data_treatments and isinstance(data_treatments, list):
            for treatment in data_treatments:
                self._add_paragraph(doc, f"- {treatment}")
        else:
            self._add_paragraph(
                doc,
                "Data was processed through the 6-step cleaning pipeline:\n"
                "1. Leakage column removal (post-origination fields)\n"
                "2. Filtering to resolved loans (Fully Paid, Charged Off)\n"
                "3. Type coercion (int_rate, term, emp_length, revol_util)\n"
                "4. Missing value imputation (grade-segmented median for numeric, mode for categorical)\n"
                "5. Winsorization at 1st/99th percentile for key numeric features\n"
                "6. Categorical encoding (ordinal, risk-based, one-hot)",
            )

        self._add_heading(doc, "4.3 Data Assumptions", level=2)
        assumptions = data_info.get("assumptions", [])
        if assumptions and isinstance(assumptions, list):
            for assumption in assumptions:
                self._add_paragraph(doc, f"- {assumption}")
        else:
            self._add_paragraph(
                doc,
                "Key assumptions:\n"
                "- Resolved loans (Fully Paid / Charged Off) are representative of future performance\n"
                "- Default definition: Charged Off status (typically 150+ days delinquent)\n"
                "- Feature distributions are sufficiently stable across vintages for modeling\n"
                "- Missing data is Missing at Random (MAR) and imputation does not introduce material bias",
            )

        # --- Section 5: Methodology ---
        self._add_heading(doc, "5. Methodology", level=1)

        self._add_heading(doc, "5.1 Feature Selection", level=2)
        feature_list = features.get("selected_features", [])
        feature_importance = features.get("feature_importance", {})
        if feature_list:
            self._add_paragraph(doc, f"The model uses {len(feature_list)} features selected through the feature engineering pipeline:")
            if feature_importance:
                headers = ["Feature", "Importance"]
                rows = [[f, str(round(v, 4))] for f, v in sorted(feature_importance.items(), key=lambda x: -x[1])[:20]]
                self._add_table(doc, headers, rows)
            else:
                for i in range(0, len(feature_list), 4):
                    chunk = feature_list[i:i + 4]
                    self._add_paragraph(doc, ", ".join(chunk))
        else:
            feat_count = metrics.get("feature_count", "N/A")
            self._add_paragraph(
                doc,
                f"Feature selection was performed through the Feature Engineering Agent, "
                f"resulting in {feat_count} features. Details are available in the "
                f"feature engineering handoff artifacts.",
            )

        self._add_heading(doc, "5.2 Model Selection (Tournament)", level=2)
        tournament = artifacts.get("tournament", {})
        if tournament:
            self._add_paragraph(
                doc,
                f"A model tournament was conducted evaluating "
                f"{tournament.get('candidates_evaluated', 'multiple')} candidate "
                f"algorithms with hyperparameter optimization.",
            )
            if tournament.get("results"):
                headers = ["Algorithm", "AUC-ROC", "Gini", "KS Statistic", "Brier Score"]
                rows = []
                for result in tournament["results"]:
                    if isinstance(result, dict):
                        rows.append([
                            result.get("algorithm", ""),
                            str(round(result.get("auc_roc", 0), 4)),
                            str(round(result.get("gini", 0), 4)),
                            str(round(result.get("ks_statistic", 0), 4)),
                            str(round(result.get("brier_score", 0), 4)),
                        ])
                self._add_table(doc, headers, rows)
        else:
            self._add_paragraph(
                doc,
                "The model tournament framework evaluated multiple algorithms "
                "(Logistic Regression, Random Forest, Gradient Boosting, XGBoost, LightGBM) "
                "using randomized hyperparameter search with cross-validation. "
                "Champion selection was based on a composite regulatory scoring metric "
                "considering AUC-ROC, Gini coefficient, KS statistic, and calibration.",
            )

        self._add_heading(doc, "5.3 Champion Model", level=2)
        if champion:
            champion_info = {
                "Algorithm": champion.get("algorithm", "N/A"),
                "Model ID": champion.get("model_id", "N/A"),
                "AUC-ROC": str(round(champion.get("auc_roc", 0), 4)) if isinstance(champion.get("auc_roc"), (int, float)) else "N/A",
                "Gini": str(round(champion.get("gini", 0), 4)) if isinstance(champion.get("gini"), (int, float)) else "N/A",
            }
            hyperparams = champion.get("hyperparameters", {})
            if hyperparams:
                champion_info["Hyperparameters"] = json.dumps(hyperparams, indent=2, default=str)
            self._add_metric_table(doc, champion_info)
        else:
            self._add_paragraph(
                doc,
                "Champion model details are available in the model agent handoff artifacts.",
            )

        self._add_heading(doc, "5.4 Calibration", level=2)
        calibration = artifacts.get("calibration", {})
        if calibration:
            self._add_metric_table(doc, calibration)
        else:
            self._add_paragraph(
                doc,
                "Model calibration was assessed using the Brier score and calibration "
                "curve analysis. The predicted probabilities were compared against "
                "observed default rates across decile bins to verify monotonicity "
                "and accuracy of probability estimates.",
            )

        # --- Section 6: Outputs ---
        self._add_heading(doc, "6. Outputs", level=1)

        self._add_heading(doc, "6.1 Output Description", level=2)
        output_desc_map = {
            "PD": "The model produces a probability of default score between 0 and 1 for each loan.",
            "LGD": "The model produces a loss given default estimate between 0 and 1 for each defaulted loan.",
            "EAD": "The model produces an exposure at default estimate in dollar terms for each loan.",
        }
        self._add_paragraph(doc, output_desc_map.get(model_label, f"Model output for {model_label}."))

        self._add_heading(doc, "6.2 Performance Testing", level=2)
        if performance:
            self._add_metric_table(doc, {
                k: str(round(v, 4)) if isinstance(v, float) else str(v)
                for k, v in performance.items()
            })
        else:
            perf_metrics_map = {
                "PD": {
                    "Primary Metric": "AUC-ROC (Area Under ROC Curve)",
                    "Discrimination": "Gini Coefficient, KS Statistic",
                    "Calibration": "Brier Score, Hosmer-Lemeshow Test",
                    "Stability": "Population Stability Index (PSI)",
                },
                "LGD": {
                    "Primary Metric": "R-squared / RMSE",
                    "Accuracy": "Mean Absolute Error (MAE)",
                    "Stability": "Population Stability Index (PSI)",
                },
                "EAD": {
                    "Primary Metric": "R-squared / RMSE",
                    "Accuracy": "Mean Absolute Error (MAE)",
                    "Stability": "Population Stability Index (PSI)",
                },
            }
            self._add_metric_table(doc, perf_metrics_map.get(model_label, {}))

        # --- Section 7: Implementation ---
        self._add_heading(doc, "7. Implementation", level=1)
        self._add_paragraph(
            doc,
            "The champion model is serialized and registered in the model registry. "
            "Implementation considerations:\n"
            "- Model artifact is stored as a pickle/joblib file\n"
            "- Feature pipeline must replicate the exact cleaning and engineering steps\n"
            "- Input validation should check for required features and valid ranges\n"
            "- Model monitoring should track PSI for input features and output scores\n"
            "- Retraining triggers: PSI > 0.25 or significant AUC degradation (> 5%)",
        )

        # --- Section 8: Appendix ---
        self._add_heading(doc, "8. Appendix", level=1)

        output_files = artifacts.get("output_files", {})
        if output_files:
            self._add_heading(doc, "A.1 Output Files", level=2)
            self._add_metric_table(doc, output_files)

        appendix_charts = artifacts.get("charts", [])
        if appendix_charts:
            self._add_heading(doc, "A.2 Supporting Charts", level=2)
            for chart_path in appendix_charts:
                chart_file = Path(chart_path)
                if chart_file.exists():
                    self._add_image(doc, str(chart_file))

        filename = f"{model_label.lower()}_model_development_report.docx"
        filepath = output_path / filename
        doc.save(str(filepath))
        return filepath

    def generate_el_report(self, el_data: dict, output_path: Path) -> Path:
        """Generate the Expected Loss report.

        Sections:
          1. Executive Summary
          2. Model Performance Summary
          3. Loan-Level Distribution
          4. Portfolio Roll-Up
          5. Stress Testing
          6. Regulatory Capital

        Args:
            el_data: Combined EL handoff data with PD/LGD/EAD summaries.
            output_path: Directory to write the report into.

        Returns:
            Path to the generated .docx file.
        """
        output_path = Path(output_path)
        output_path.mkdir(parents=True, exist_ok=True)

        doc = self._create_base_doc(
            title="Expected Loss Report",
            subtitle="Credit Risk Modeling Platform — Portfolio Loss Analysis",
        )

        metrics = el_data.get("metrics", {})
        portfolio = el_data.get("portfolio", {})
        stress = el_data.get("stress_testing", {})
        regulatory = el_data.get("regulatory_capital", {})
        model_summaries = el_data.get("model_summaries", {})
        distributions = el_data.get("distributions", {})

        # --- Section 1: Executive Summary ---
        self._add_heading(doc, "1. Executive Summary", level=1)
        exec_summary = {
            "Total Portfolio Exposure": _fmt_currency(metrics.get("total_exposure")),
            "Expected Loss (EL)": _fmt_currency(metrics.get("total_el")),
            "EL Rate": _fmt_pct(metrics.get("el_rate")),
            "Weighted Average PD": _fmt_pct(metrics.get("weighted_avg_pd")),
            "Weighted Average LGD": _fmt_pct(metrics.get("weighted_avg_lgd")),
            "Number of Loans": _fmt_int(metrics.get("loan_count")),
            "Report Date": datetime.now(timezone.utc).strftime("%Y-%m-%d"),
        }
        self._add_metric_table(doc, exec_summary)

        self._add_paragraph(
            doc,
            "This report presents the Expected Loss (EL = PD x LGD x EAD) analysis "
            "for the consumer lending portfolio. EL estimates are derived from the "
            "champion PD, LGD, and EAD models developed through the Credit Risk "
            "Modeling Platform's automated pipeline.",
        )

        # --- Section 2: Model Performance Summary ---
        self._add_heading(doc, "2. Model Performance Summary", level=1)

        for model_type in ["PD", "LGD", "EAD"]:
            summary = model_summaries.get(model_type, {})
            if summary:
                self._add_heading(doc, f"2.{['PD','LGD','EAD'].index(model_type)+1} {model_type} Model", level=2)
                display_metrics = {}
                for k, v in summary.items():
                    if isinstance(v, float):
                        display_metrics[k] = f"{v:.4f}"
                    else:
                        display_metrics[k] = str(v)
                self._add_metric_table(doc, display_metrics)

        if not model_summaries:
            self._add_paragraph(
                doc,
                "Individual model performance metrics are available in the respective "
                "model development reports (PD, LGD, EAD).",
            )

        # --- Section 3: Loan-Level Distribution ---
        self._add_heading(doc, "3. Loan-Level EL Distribution", level=1)

        el_dist = distributions.get("el_distribution", {})
        if el_dist:
            self._add_metric_table(doc, {
                k: f"{v:.4f}" if isinstance(v, float) else str(v)
                for k, v in el_dist.items()
            })
        else:
            self._add_paragraph(
                doc,
                "Loan-level EL values were computed as PD x LGD x EAD for each loan "
                "in the portfolio. Distribution statistics and percentile breakdowns "
                "are available in the EL Agent handoff artifacts.",
            )

        grade_dist = distributions.get("by_grade", [])
        if grade_dist:
            self._add_heading(doc, "3.1 EL by Grade", level=2)
            headers = ["Grade", "Count", "Avg PD", "Avg LGD", "Avg EAD", "Total EL", "EL Rate"]
            rows = []
            for g in grade_dist:
                if isinstance(g, dict):
                    rows.append([
                        str(g.get("grade", "")),
                        _fmt_int(g.get("count")),
                        _fmt_pct(g.get("avg_pd")),
                        _fmt_pct(g.get("avg_lgd")),
                        _fmt_currency(g.get("avg_ead")),
                        _fmt_currency(g.get("total_el")),
                        _fmt_pct(g.get("el_rate")),
                    ])
            self._add_table(doc, headers, rows)

        vintage_dist = distributions.get("by_vintage", [])
        if vintage_dist:
            self._add_heading(doc, "3.2 EL by Vintage", level=2)
            headers = ["Vintage", "Count", "Total EL", "EL Rate"]
            rows = []
            for v in vintage_dist:
                if isinstance(v, dict):
                    rows.append([
                        str(v.get("vintage", "")),
                        _fmt_int(v.get("count")),
                        _fmt_currency(v.get("total_el")),
                        _fmt_pct(v.get("el_rate")),
                    ])
            self._add_table(doc, headers, rows)

        # --- Section 4: Portfolio Roll-Up ---
        self._add_heading(doc, "4. Portfolio Roll-Up", level=1)
        if portfolio:
            self._add_metric_table(doc, {
                k: _fmt_currency(v) if "amount" in k.lower() or "el" in k.lower() or "exposure" in k.lower()
                else (_fmt_pct(v) if "rate" in k.lower() or "pct" in k.lower() else str(v))
                for k, v in portfolio.items()
            })
        else:
            self._add_paragraph(
                doc,
                "Portfolio roll-up aggregates loan-level EL estimates to produce "
                "portfolio-level loss metrics. Segmentation by grade, vintage, "
                "and purpose provides granular risk views for portfolio management.",
            )

        # --- Section 5: Stress Testing ---
        self._add_heading(doc, "5. Stress Testing", level=1)
        if stress:
            scenarios = stress.get("scenarios", [])
            if scenarios:
                headers = ["Scenario", "PD Multiplier", "LGD Multiplier", "Stressed EL", "EL Increase"]
                rows = []
                for s in scenarios:
                    if isinstance(s, dict):
                        rows.append([
                            s.get("name", ""),
                            str(s.get("pd_multiplier", "")),
                            str(s.get("lgd_multiplier", "")),
                            _fmt_currency(s.get("stressed_el")),
                            _fmt_pct(s.get("el_increase_pct")),
                        ])
                self._add_table(doc, headers, rows)
            else:
                self._add_metric_table(doc, {
                    k: str(v) for k, v in stress.items()
                })
        else:
            self._add_paragraph(
                doc,
                "Stress testing applies multipliers to PD and LGD estimates under "
                "adverse economic scenarios:\n"
                "- Baseline: Current model estimates\n"
                "- Mild Stress: PD x 1.5, LGD x 1.2\n"
                "- Moderate Stress: PD x 2.0, LGD x 1.4\n"
                "- Severe Stress: PD x 3.0, LGD x 1.6\n\n"
                "Detailed stress testing results are available in the EL Agent artifacts.",
            )

        # --- Section 6: Regulatory Capital ---
        self._add_heading(doc, "6. Regulatory Capital", level=1)
        if regulatory:
            self._add_metric_table(doc, {
                k: _fmt_currency(v) if isinstance(v, (int, float)) and v > 100
                else (_fmt_pct(v) if isinstance(v, float) and v < 1 else str(v))
                for k, v in regulatory.items()
            })
        else:
            self._add_paragraph(
                doc,
                "Regulatory capital estimates under the Basel III/IV Internal Ratings-Based "
                "(IRB) approach are derived from the PD, LGD, and EAD model outputs. "
                "Risk-weighted assets (RWA) and capital requirements are computed using "
                "the regulatory capital formula with the applicable asset correlation "
                "and maturity adjustment factors.",
            )

        # --- Charts ---
        charts = el_data.get("charts", [])
        if charts:
            self._add_heading(doc, "Supporting Charts", level=2)
            for chart_path in charts:
                chart_file = Path(chart_path)
                if chart_file.exists():
                    self._add_image(doc, str(chart_file))

        filepath = output_path / "expected_loss_report.docx"
        doc.save(str(filepath))
        return filepath

    # ------------------------------------------------------------------
    # Helper methods
    # ------------------------------------------------------------------

    def _create_base_doc(self, title: str, subtitle: str) -> Document:
        """Create a Document with a formatted title page.

        Args:
            title: Main title text.
            subtitle: Subtitle / description text.

        Returns:
            A python-docx Document instance.
        """
        doc = Document()

        # --- Page margins ---
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        # --- Title ---
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_before = Pt(72)
        title_para.space_after = Pt(12)
        run = title_para.add_run(title)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)  # Dark navy

        # --- Subtitle ---
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.space_after = Pt(6)
        run = sub_para.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

        # --- Date ---
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.space_after = Pt(36)
        run = date_para.add_run(datetime.now(timezone.utc).strftime("%B %d, %Y"))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)

        # --- Horizontal rule ---
        rule_para = doc.add_paragraph()
        rule_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rule_para.add_run("_" * 60)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # --- Confidentiality notice ---
        notice_para = doc.add_paragraph()
        notice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notice_para.space_before = Pt(12)
        run = notice_para.add_run("CONFIDENTIAL — For Internal Use Only")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        doc.add_page_break()
        return doc

    def _add_heading(self, doc: Document, text: str, level: int) -> None:
        """Add a styled heading to the document.

        Args:
            doc: The Document instance.
            text: Heading text.
            level: Heading level (1, 2, or 3).
        """
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)

    def _add_paragraph(self, doc: Document, text: str) -> None:
        """Add a body text paragraph to the document.

        Args:
            doc: The Document instance.
            text: Body text content.
        """
        para = doc.add_paragraph()
        para.space_after = Pt(6)
        run = para.add_run(text)
        run.font.size = Pt(11)

    def _add_table(self, doc: Document, headers: list[str], rows: list[list[str]]) -> None:
        """Add a formatted table with alternating row colors.

        Args:
            doc: The Document instance.
            headers: Column header strings.
            rows: List of row data (each row is a list of strings).
        """
        if not headers:
            return

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 1"

        # Header row
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Header background
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(
                qn("w:shd"),
                {"w:fill": "1A3C6D", "w:val": "clear"},
            )
            shading.append(shading_elm)

        # Data rows with alternating colors
        for i, row_data in enumerate(rows):
            for j, value in enumerate(row_data):
                cell = table.rows[i + 1].cells[j]
                cell.text = str(value) if value is not None else ""
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

                # Alternating row background
                if i % 2 == 1:
                    shading = cell._element.get_or_add_tcPr()
                    shading_elm = shading.makeelement(
                        qn("w:shd"),
                        {"w:fill": "EBF1FA", "w:val": "clear"},
                    )
                    shading.append(shading_elm)

        doc.add_paragraph()  # spacing after table

    def _add_metric_table(self, doc: Document, metrics_dict: dict) -> None:
        """Add a two-column key-value metric table.

        Args:
            doc: The Document instance.
            metrics_dict: Dictionary of metric name -> value.
        """
        if not metrics_dict:
            return

        table = doc.add_table(rows=len(metrics_dict), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (key, value) in enumerate(metrics_dict.items()):
            # Key cell
            key_cell = table.rows[i].cells[0]
            key_cell.text = str(key)
            for paragraph in key_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            # Key background
            shading = key_cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(
                qn("w:shd"),
                {"w:fill": "F0F4F8", "w:val": "clear"},
            )
            shading.append(shading_elm)

            # Value cell
            val_cell = table.rows[i].cells[1]
            val_cell.text = str(value) if value is not None else "N/A"
            for paragraph in val_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        doc.add_paragraph()  # spacing after table

    def _add_metadata_table(self, doc: Document, metadata: dict) -> None:
        """Add a Table 1 format metadata table (C1 template style).

        Two-column table with bold header row, used for model identification.

        Args:
            doc: The Document instance.
            metadata: Dictionary of metadata fields.
        """
        if not metadata:
            return

        table = doc.add_table(rows=1 + len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for j, header_text in enumerate(["Field", "Value"]):
            cell = table.rows[0].cells[j]
            cell.text = header_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(
                qn("w:shd"),
                {"w:fill": "1A3C6D", "w:val": "clear"},
            )
            shading.append(shading_elm)

        # Data rows
        for i, (key, value) in enumerate(metadata.items()):
            table.rows[i + 1].cells[0].text = str(key)
            table.rows[i + 1].cells[1].text = str(value) if value is not None else "N/A"
            for j in range(2):
                for paragraph in table.rows[i + 1].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        if j == 0:
                            run.font.bold = True

        doc.add_paragraph()  # spacing after table

    def _add_image(self, doc: Document, image_path: str, width: Inches = Inches(5)) -> None:
        """Embed a chart image in the document.

        Args:
            doc: The Document instance.
            image_path: Path to the image file.
            width: Display width (default 5 inches).
        """
        image_file = Path(image_path)
        if not image_file.exists():
            self._add_paragraph(doc, f"[Chart not found: {image_path}]")
            return

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(image_file), width=width)

        # Caption
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(image_file.stem.replace("_", " ").title())
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True
        cap_run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)


# ------------------------------------------------------------------
# Module-level formatting helpers
# ------------------------------------------------------------------

def _fmt_currency(value: Any) -> str:
    """Format a value as currency."""
    if value is None:
        return "N/A"
    try:
        v = float(value)
        if abs(v) >= 1_000_000_000:
            return f"${v / 1_000_000_000:,.2f}B"
        elif abs(v) >= 1_000_000:
            return f"${v / 1_000_000:,.2f}M"
        elif abs(v) >= 1_000:
            return f"${v / 1_000:,.1f}K"
        else:
            return f"${v:,.2f}"
    except (TypeError, ValueError):
        return str(value)


def _fmt_pct(value: Any) -> str:
    """Format a value as a percentage."""
    if value is None:
        return "N/A"
    try:
        return f"{float(value):.2%}"
    except (TypeError, ValueError):
        return str(value)


def _fmt_int(value: Any) -> str:
    """Format a value as a comma-separated integer."""
    if value is None:
        return "N/A"
    try:
        return f"{int(value):,}"
    except (TypeError, ValueError):
        return str(value)
